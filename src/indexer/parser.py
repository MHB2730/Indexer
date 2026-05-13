"""Text extraction from PDF and DOCX, with OCR fallback for scans."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

from .ocr import MIN_TEXT_LEN_PER_PAGE, is_available as ocr_available, ocr_page


# Cap how much content we read per file when building the matching corpus
# — full document text matters but we don't need megabytes of it.
MATCH_TEXT_CHAR_CAP = 30_000


@dataclass
class ExtractedDoc:
    path: Path
    text: str
    pages: list[str]


def extract(path: Path, use_ocr: bool = True) -> ExtractedDoc:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path, use_ocr=use_ocr)
    if suffix in {".docx", ".doc"}:
        return _extract_docx(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def _extract_pdf(path: Path, use_ocr: bool) -> ExtractedDoc:
    pages: list[str] = []
    do_ocr = use_ocr and ocr_available()
    with fitz.open(path) as doc:
        for page in doc:
            text = page.get_text("text") or ""
            if do_ocr and len(text.strip()) < MIN_TEXT_LEN_PER_PAGE:
                ocr_text = ocr_page(page)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
            pages.append(text)
    return ExtractedDoc(path=path, text="\n".join(pages), pages=pages)


def _extract_docx(path: Path) -> ExtractedDoc:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    text = "\n".join(parts)
    return ExtractedDoc(path=path, text=text, pages=[text])


def first_page_text(path: Path, max_chars: int = 4000) -> str:
    """Cheap fingerprint: first page (or first N chars for DOCX)."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        with fitz.open(path) as doc:
            if len(doc) == 0:
                return ""
            text = (doc[0].get_text("text") or "")
            if len(text.strip()) < MIN_TEXT_LEN_PER_PAGE and ocr_available():
                ocr_text = ocr_page(doc[0])
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
            return text[:max_chars]
    if suffix in {".docx", ".doc"}:
        return _extract_docx(path).text[:max_chars]
    return ""


def matching_text(path: Path) -> str:
    """Full text for matching purposes — capped, OCR-augmented if available."""
    try:
        doc = extract(path, use_ocr=True)
    except Exception:
        return ""
    return doc.text[:MATCH_TEXT_CHAR_CAP]
